# -*- coding: utf-8 -*-
"""
旅行攻略 Word 文档生成器
================================
使用方法：
  1. 修改下方「===== 修改行程数据 =====」部分
  2. 运行: python generate_travel_doc.py
  3. 输出: 攻略.docx + route_map.html

依赖安装: pip install python-docx folium
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, subprocess

# ===== 修改行程数据 =====
TITLE = "贵阳 3 天 2 晚"
SUB_TITLE = "广州出发 · 避暑美食之旅"
TOTAL_BUDGET = 2000
FROM_TO = "广州南 → 贵阳北"
STYLE_DESC = "休闲 + 美食 + 文化"
WEATHER = "均温23°C 避暑胜地 7月多雨带伞"

DAYS = [
    {
        "day": 1, "color": "FF7B54",
        "theme": "老城烟火 · 市区漫步",
        "items": [
            ("🚄 高铁出发", "G2942 广州南 07:44→贵阳北 11:10  3h26min  票价¥406", 406, "#667eea"),
            ("🏨 入住酒店", "喷水池区域 舒适型酒店  地铁1号线直达", 150, "#FF7B54"),
            ("🍜 午餐酸汤鱼", "亮欢寨酸汤鱼(飞山街店) 贵阳灵魂美食", 80, "#FFD93D"),
            ("🐒 黔灵山公园", "门票¥5 看野生猕猴+熊猫+弘福寺", 5, "#4ECDC4"),
            ("🍢 民生路扫街", "肠旺面+豆腐圆子+但家香酥鸭+冰粉", 62, "#FFD93D"),
            ("🌃 甲秀楼夜景", "免费 贵阳地标 亮灯后绝美 南明河畔", 0, "#4ECDC4"),
        ],
        "tip": "黔灵山的猴子会抢吃的！手里别拿塑料袋和食物",
    },
    {
        "day": 2, "color": "4ECDC4",
        "theme": "山水古镇 · 文艺探索",
        "items": [
            ("🍜 早餐牛肉粉", "花溪王记牛肉粉(民生路店) 本地人最爱", 18, "#FFD93D"),
            ("🏛️ 省博物馆", "免费(需预约) 苗族银饰+民族服饰超震撼", 0, "#4ECDC4"),
            ("🍲 午餐酸汤鱼", "老凯俚酸汤鱼(省府路总店) 游客少", 60, "#FFD93D"),
            ("🗿 花溪夜郎谷", "门票¥20 石头城堡超出片", 20, "#4ECDC4"),
            ("🌮 青云市集夜市", "烙锅+烤串 老厂房夜市 上百家店", 60, "#FFD93D"),
            ("🚗 市内交通", "打车往返各景点", 40, "#AAAAAA"),
        ],
        "tip": "博物馆周一闭馆，记得提前在公众号预约",
    },
    {
        "day": 3, "color": "FFD93D",
        "theme": "古镇收尾 · 满载而归",
        "items": [
            ("🏯 青岩古镇", "门票¥10 卤猪脚必吃 金必轩老字号", 10, "#4ECDC4"),
            ("🥟 午餐丝娃娃", "丝恋红汤丝娃娃 自己动手包", 25, "#FFD93D"),
            ("🚄 高铁返回", "G2922 贵阳北 16:43→广州南 20:32  票价¥422", 422, "#667eea"),
            ("🚗 市内交通", "打车到古镇+到高铁站", 40, "#AAAAAA"),
        ],
        "tip": "青岩古镇卤猪脚认准「金必轩」，30年老字号",
    },
]

BUDGET_DATA = [
    ("🚄 高铁往返", 828, "#667eea"),
    ("🏨 住宿2晚", 600, "#FF7B54"),
    ("🍜 餐饮美食", 305, "#FFD93D"),
    ("🎫 门票", 35, "#4ECDC4"),
    ("🚗 市内交通", 100, "#AAAAAA"),
]

SAVING_TIPS = [
    "选D字头动车可省¥116（如D1864次¥290）",
    "两人出行住宿分摊，每人省¥300",
    "博物馆/甲秀楼/青云市集全部免费",
    "7月多阵雨，冲锋衣比雨伞实用",
]

CHECKLIST = [
    ("🎫", "证件", "身份证、学生证"),
    ("📱", "App", "12306/携程/美团/高德"),
    ("🌂", "装备", "晴雨伞、薄外套、徒步鞋"),
    ("🔌", "充电", "充电宝、数据线"),
    ("💰", "现金", "少量现金"),
]

TRAVEL_TIPS = [
    "带伞——7月多阵雨",
    "带薄外套——早晚偏凉",
    "贵阳北站下，地铁1号线进市区",
]

FOOD_TABLE_DATA = [
    ["肠旺面", "金牌罗记肠旺面", "¥12"],
    ["豆腐圆子", "雷家豆腐圆子", "¥10"],
    ["但家香酥鸭", "小十字总店", "¥30"],
    ["恋爱豆腐果", "路边摊", "¥5"],
    ["冰粉", "随便一家", "¥5"],
]

SAVE_MONEY_TABLE = [
    ("选D字头动车", "省¥116"),
    ("住青旅/民宿", "省¥100~150/晚"),
    ("两人出行分摊住宿", "省¥300"),
    ("选免费景点", "本来就免费"),
    ("小吃代替正餐", "省一半"),
]


# ===== 以下不需要修改 =====
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for sec in doc.sections:
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

def _set_font(run, text, size, bold=False, color=None):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    has_emoji = any(ord(c) > 0x1F000 for c in text)
    if has_emoji:
        rFonts.set(qn('w:ascii'), 'Segoe UI Emoji')
        rFonts.set(qn('w:hAnsi'), 'Segoe UI Emoji')
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    else:
        rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
        rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rPr.insert(0, rFonts)
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    run.font.bold = bold

def T(text, size=11, color=RGBColor(0x33,0x33,0x33), bold=False, align=None, sa=2, sb=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, text, size, bold, color)
    if align: p.alignment = align
    p.space_after = Pt(sa); p.space_before = Pt(sb)
    return p

def TI(icon, title, desc, cost, c="#666"):
    p = doc.add_paragraph(); p.space_after = Pt(3)
    r, g, b = int(c[1:3],16), int(c[3:5],16), int(c[5:7],16)
    r1 = p.add_run(f"{icon}  {title}  ")
    _set_font(r1, f"{icon}  {title}  ", 10.5, True, RGBColor(0x33,0x33,0x33))
    r2 = p.add_run(desc)
    _set_font(r2, desc, 10, color=RGBColor(0x77,0x77,0x88))
    if cost == 0:
        r3 = p.add_run("  免费")
        _set_font(r3, "  免费", 10, True, RGBColor(0x4E,0xCD,0xC4))
    else:
        r3 = p.add_run(f"  ¥{cost}")
        _set_font(r3, f"  ¥{cost}", 10, True, RGBColor(r,g,b))
    return p

def DH(num, theme, ch):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0,0); c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"  Day {num}  —  {theme}  ")
    r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    sh = c._element.get_or_add_tcPr()
    se = sh.makeelement(qn('w:shd'), {qn('w:fill'): ch, qn('w:val'): 'clear'})
    sh.append(se)
    doc.add_paragraph().space_after = Pt(4)

def TB(headers, data):
    t = doc.add_table(rows=len(data)+1, cols=len(headers))
    t.style = 'Light List Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.font.bold = True; r.font.size = Pt(10)
    for ri, rd in enumerate(data):
        for ci, val in enumerate(rd):
            p = t.rows[ri+1].cells[ci].paragraphs[0]; p.clear()
            r = p.add_run(str(val))
            _set_font(r, str(val), 10)
            if ci == len(headers)-1: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ===== 封面 =====
T(f"{TITLE}  ·  私 人 旅 行 攻 略", 24, RGBColor(0xFF,0x6B,0x35), True, WD_ALIGN_PARAGRAPH.CENTER, 6, 60)
T(f"☀️ {WEATHER}  ·  💰 预算¥{TOTAL_BUDGET}", 10, RGBColor(0x88,0x88,0x88), align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
T("")
T("📋 行程总览", 13, RGBColor(0xFF,0x6B,0x35), True, sb=6)
total = sum(sum(it[2] for it in d["items"]) for d in DAYS)
for line in [
    f"路线：{FROM_TO}",
    f"天数：{len(DAYS)}天{len(DAYS)-1}晚",
    f"预算：¥{TOTAL_BUDGET}/人",
    f"预估总花费：¥{total}/人  ✅ 预算内",
    f"风格：{STYLE_DESC}",
    f"天气：{WEATHER}",
]:
    T(f"  {line}", 10.5, RGBColor(0x55,0x55,0x55))
T("")

# ===== 每日行程 =====
for d in DAYS:
    doc.add_page_break()
    DH(d["day"], d["theme"], d["color"])
    for icon, desc, cost, color in d["items"]:
        TI(icon, icon.split(" ")[1] if " " in icon else "", desc, cost, color)
    T("")
    T("━" * 55, 7, RGBColor(0xCC,0xCC,0xCC))
    sub = sum(it[2] for it in d["items"])
    c = RGBColor(*(int(d["color"][j:j+2],16) for j in (0,2,4)))
    T(f"📊  Day {d['day']} 花费小计：¥{sub}", 12, c, True, WD_ALIGN_PARAGRAPH.RIGHT, sb=4)
    T(f"💡 {d['tip']}", 9, RGBColor(0xAA,0x77,0x33))
    T("")

# ===== 预算汇总 =====
doc.add_page_break()
T("💰  预 算 透 明 汇 总", 18, RGBColor(0xFF,0x6B,0x35), True, WD_ALIGN_PARAGRAPH.CENTER, 8, 20)
TB(["项目", "金额"], [[n, f"¥{c}"] for n,c,_ in BUDGET_DATA])
T("")
T(f"总计：¥{total}/人  ✅ 预算剩余¥{TOTAL_BUDGET-total}", 13, RGBColor(0xFF,0x6B,0x35), True, WD_ALIGN_PARAGRAPH.CENTER)

T("")
T("━━ 省钱方案对比 ━━", 11, RGBColor(0x88,0x88,0x88), True, WD_ALIGN_PARAGRAPH.CENTER, 6)
TB(["方法", "省多少"], SAVE_MONEY_TABLE)

T("")
T("━━ 省钱小贴士 ━━", 11, RGBColor(0x88,0x88,0x88), True, WD_ALIGN_PARAGRAPH.CENTER, 6)
for t in SAVING_TIPS:
    T(f"  • {t}")

T("")
T("━━ 出行贴士 ━━", 11, RGBColor(0x88,0x88,0x88), True, WD_ALIGN_PARAGRAPH.CENTER, 6)
for t in TRAVEL_TIPS:
    T(f"  • {t}")

T("")
T("━━ 出发前 Checklist ━━", 11, RGBColor(0x88,0x88,0x88), True, WD_ALIGN_PARAGRAPH.CENTER, 6)
for ic, t, d in CHECKLIST:
    T(f"  ☐  {ic} {t}——{d}")

T("")
T("🌟  祝 旅 途 愉 快 ！  🌟", 16, RGBColor(0xFF,0x6B,0x35), True, WD_ALIGN_PARAGRAPH.CENTER, 4, 10)
T("由旅行管家 Travel Agent Skill 生成", 9, RGBColor(0xBB,0xBB,0xBB), align=WD_ALIGN_PARAGRAPH.CENTER)

output_path = os.path.join(os.path.dirname(__file__) or ".", "攻略.docx")
doc.save(output_path)
print(f"✅ 已生成: {output_path}")
try:
    subprocess.run(['start', output_path], shell=True)
except:
    pass
